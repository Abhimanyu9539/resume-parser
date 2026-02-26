"""
Document text extraction service for PDF and DOCX files.
"""

import sys
import pymupdf  # PyMuPDF
from docx import Document
from app.utils.custom_exception import CustomException
from app.utils.logger import logger


class DocumentExtractor:
    """Extract text content from PDF and DOCX files."""

    def extract_from_pdf(file_path: str) -> str:
        """
        Extract text from a PDF file.

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text content
        """
        try:
            logger.info(f"Extracting text from PDF: {file_path}")
            text = ""

            # Open the PDF file
            doc = pymupdf.open(file_path)

            # Extract text from each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                text += page.get_text()

            doc.close()

            if not text.strip():
                logger.warning(f"No text extracted from PDF: {file_path}")
                raise ValueError(
                    "PDF file appears to be empty or contains no extractable text"
                )

            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text.strip()

        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {str(e)}")
            raise CustomException(e, sys)

    def extract_from_docx(file_path: str) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content
        """
        try:
            logger.info(f"Extracting text from DOCX: {file_path}")
            text = ""

            # Open the DOCX file
            doc = Document(file_path)

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            if not text.strip():
                logger.warning(f"No text extracted from DOCX: {file_path}")
                raise ValueError(
                    "DOCX file appears to be empty or contains no extractable text"
                )

            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text.strip()

        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {str(e)}")
            raise CustomException(e, sys)

    def extract_text(file_path: str, file_extension: str) -> str:
        """
        Extract text from a file based on its extension.

        Args:
            file_path: Path to the file
            file_extension: File extension (.pdf or .docx)

        Returns:
            Extracted text content
        """
        try:
            logger.info(
                f"Extracting text from file: {file_path} (type: {file_extension})"
            )

            if file_extension.lower() == ".pdf":
                return DocumentExtractor.extract_from_pdf(file_path)
            elif file_extension.lower() in [".docx", ".doc"]:
                return DocumentExtractor.extract_from_docx(file_path)
            else:
                error_msg = f"Unsupported file format: {file_extension}. Only PDF and DOCX are supported."
                logger.error(error_msg)
                raise ValueError(error_msg)

        except Exception as e:
            logger.error(f"Document extraction failed: {str(e)}")
            raise CustomException(e, sys)
